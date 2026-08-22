"""
make_algorithm_box.py
=====================
Produit l'encadre de pseudocode de l'article (Algorithm 1) au format
Word (.docx), dans la presentation attendue par les revues Elsevier :
filet horizontal en haut, legende au-dessus, filet sous l'en-tete
Input/Output, filet de fermeture en bas, lignes numerotees et indentees.

POURQUOI CE SCRIPT PLUTOT QU'UN COPIER-COLLER DEPUIS LATEX
    L'environnement algorithm2e de la these ne se transpose pas dans
    Word : la numerotation, l'indentation et les filets sont produits par
    le paquet, pas par le texte. Ce script les reconstruit explicitement.

DEUX ECARTS PAR RAPPORT AU PSEUDOCODE DE LA THESE — voir le rapport
imprime en fin d'execution. Les lignes sont editables dans LINES
ci-dessous si vous souhaitez revenir a la version d'origine.

Usage :
    python make_algorithm_box.py --outdir tables_article
"""

import argparse
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERREUR : python-docx n'est pas installe.")
    print("Lancez : pip install python-docx")
    sys.exit(1)

CAPTION = "Algorithm 1. Pseudo-code of ABC-SHAP"

# En-tete Input / Output — pas de numero de ligne, separe par un filet.
HEADER = [
    ("Input:",  "f, x, D, N, SN, T_max, T_min, patience, \u03b5_\u03c6, \u03b1_mut"),
    ("Output:", "\u03c6\u0302 \u2208 \u211d\u1d39"),
]

# Corps : (niveau d'indentation, texte, type)
#   type = 'code'    ligne numerotee
#          'comment' commentaire, non numerote, en italique
LINES = [
    (0, "Phase 1 \u2014 Budget allocation and initialization", 'comment'),
    (0, "N_inner \u2190 max(2(M+1), 4M)", 'code'),
    (0, "(SN_eff, T_eff) \u2190 BudgetRule(N, M, SN, T_max)", 'code'),
    (0, "N_explor \u2190 (1 + 2\u00b7T_eff)\u00b7SN_eff\u00b7N_inner ;  eval \u2190 0 ;  limit \u2190 SN_eff\u00b7M", 'code'),
    (0, "for i = 1 to SN_eff do", 'code'),
    (1, "\u03c0\u1d62 ~ Dir(1)", 'code'),
    (1, "(J\u1d62, \u03c6\u1d62) \u2190 EvalPolicy(\u03c0\u1d62, N_inner) ;  eval \u2190 eval + N_inner", 'code'),
    (0, "end for", 'code'),
    (0, "i* \u2190 argmax\u1d62 J\u1d62 ;  \u03c6_cur \u2190 \u03c6\u1d62* ;  \u03c6_prev \u2190 \u03c6_cur", 'code'),
    (0, "Phase 2 \u2014 ABC search", 'comment'),
    (0, "for t = 1 to T_eff do", 'code'),
    (1, "for i = 1 to SN_eff do", 'code'),
    (2, "\u03b4 ~ Dir(1) ;  \u03c0\u2032 \u2190 (1 \u2212 \u03b1_mut)\u00b7\u03c0\u1d62 + \u03b1_mut\u00b7\u03b4", 'code'),
    (2, "(J\u2032, \u03c6\u2032) \u2190 EvalPolicy(\u03c0\u2032, N_inner) ;  eval \u2190 eval + N_inner", 'code'),
    (2, "if J\u2032 > J\u1d62 then (\u03c0\u1d62, J\u1d62, \u03c6\u1d62) \u2190 (\u03c0\u2032, J\u2032, \u03c6\u2032) ;  trial\u1d62 \u2190 0", 'code'),
    (2, "else trial\u1d62 \u2190 trial\u1d62 + 1", 'code'),
    (1, "end for", 'code'),
    (1, "Onlooker bees: resample i with probability \u221d J\u1d62 / \u03a3J, same update", 'comment'),
    (1, "Scout bees: if trial\u1d62 > limit then reset \u03c0\u1d62 ~ Dir(1), trial\u1d62 \u2190 0", 'comment'),
    (1, "i* \u2190 argmax\u1d62 J\u1d62 ;  \u03c6_cur \u2190 \u03c6\u1d62*", 'code'),
    (1, "if t \u2265 T_min and \u2016\u03c6_cur \u2212 \u03c6_prev\u2016/\u2016\u03c6_prev\u2016 < \u03b5_\u03c6 for \u2265 patience iters then break", 'code'),
    (1, "\u03c6_prev \u2190 \u03c6_cur", 'code'),
    (0, "end for", 'code'),
    (0, "Phase 3 \u2014 Final estimation on the residual budget", 'comment'),
    (0, "\u03c0* \u2190 \u03c0\u1d62* ;  N_final \u2190 N \u2212 eval", 'code'),
    (0, "if N_final > N_inner then", 'code'),
    (1, "\u03c6\u0302 \u2190 EvalPolicy(\u03c0*, N_final)   \u25b7 fresh, disjoint coalitions", 'code'),
    (0, "else", 'code'),
    (1, "\u03c6\u0302 \u2190 \u03c6_cur              \u25b7 residual below N_inner", 'code'),
    (0, "end if", 'code'),
    (0, "return \u03c6\u0302", 'code'),
]

INDENT_CM = 0.45
CODE_FONT = 'Consolas'      # a defaut, Word substituera une police fixe
BODY_FONT = 'Times New Roman'


# ══════════════════════════════════════════════════════════════════════════
# FILETS HORIZONTAUX
# ══════════════════════════════════════════════════════════════════════════

def set_row_borders(row, top=None, bottom=None):
    """Applique un filet en haut et/ou en bas de toutes les cellules d'une
    ligne. 'sz' est en huitiemes de point : 12 = 1.5 pt, 6 = 0.75 pt."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        for el in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(el)
        b = OxmlElement('w:tcBorders')
        for side, sz in (('top', top), ('bottom', bottom)):
            e = OxmlElement('w:' + side)
            if sz:
                e.set(qn('w:val'), 'single')
                e.set(qn('w:sz'), str(sz))
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), '000000')
            else:
                e.set(qn('w:val'), 'nil')
            b.append(e)
        for side in ('left', 'right', 'insideH', 'insideV'):
            e = OxmlElement('w:' + side)
            e.set(qn('w:val'), 'nil')
            b.append(e)
        tcPr.append(b)


def fixed_layout(table, widths_cm):
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(round(Cm(w).twips))))
        grid.append(gc)
    table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
    for j, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[j].width = Cm(w)


def tight(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


# ══════════════════════════════════════════════════════════════════════════

def build(outdir):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(10)

    # ── Legende, au-dessus de l'encadre ──────────────────────────────────
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(4)
    head, rest = CAPTION.split('. ', 1)
    r = cap.add_run(head + '.')
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(' ' + rest)
    r2.italic = True
    r2.font.size = Pt(10)

    n_rows = len(HEADER) + len(LINES)
    t = doc.add_table(rows=n_rows, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    fixed_layout(t, [1.0, 14.6])

    # ── En-tete Input / Output ───────────────────────────────────────────
    for k, (lab, val) in enumerate(HEADER):
        c0, c1 = t.rows[k].cells
        c0.text = ''
        tight(c0.paragraphs[0])
        c1.text = ''
        p1 = tight(c1.paragraphs[0])
        rl = p1.add_run(lab + ' ')
        rl.bold = True
        rl.font.size = Pt(8.5)
        rl.font.name = BODY_FONT
        rv = p1.add_run(val)
        rv.font.size = Pt(8.5)
        rv.font.name = CODE_FONT

    # ── Corps ────────────────────────────────────────────────────────────
    num = 0
    for k, (lvl, text, kind) in enumerate(LINES):
        row = t.rows[len(HEADER) + k]
        c0, c1 = row.cells

        c0.text = ''
        p0 = tight(c0.paragraphs[0])
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if kind == 'code':
            num += 1
            rn = p0.add_run(f"{num}:")
            rn.font.size = Pt(8)
            rn.font.name = CODE_FONT
            rn.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        c1.text = ''
        p1 = tight(c1.paragraphs[0])
        p1.paragraph_format.left_indent = Cm(INDENT_CM * lvl)
        rt = p1.add_run(('\u25b7 ' + text) if kind == 'comment' else text)
        rt.font.size = Pt(8.5)
        rt.font.name = CODE_FONT
        if kind == 'comment':
            rt.italic = True
            rt.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    # ── Filets : haut, sous l'en-tete, bas ───────────────────────────────
    for row in t.rows:
        set_row_borders(row)
    set_row_borders(t.rows[0], top=12)
    set_row_borders(t.rows[len(HEADER) - 1], bottom=6)
    set_row_borders(t.rows[-1], bottom=12)

    os.makedirs(outdir, exist_ok=True)
    path = os.path.join(outdir, 'algorithm1_abc_shap.docx')
    doc.save(path)
    return path, num


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--outdir', default='tables_article')
    args = ap.parse_args()
    path, n = build(args.outdir)
    print(f"[algo] ecrit : {os.path.abspath(path)}")
    print(f"[algo] {n} lignes numerotees, "
          f"{sum(1 for l in LINES if l[2] == 'comment')} commentaires")
    print()
    print("ECARTS PAR RAPPORT AU PSEUDOCODE DE LA THESE — a valider :")
    print()
    print("  1. N_inner. La these ecrit  N_inner <- max(4M, N/(4*SN)),")
    print("     qui est la formule d'ORIGINE. L'article ne rapporte que la")
    print("     regle alternative, ou N_inner = max(2(M+1), 4M) et le cout")
    print("     d'exploration vaut (1 + 2*T_eff) * SN_eff * N_inner, fixe")
    print("     a l'avance. Le pseudocode a ete aligne sur la Section 3.3 ;")
    print("     sans cela il contredirait le texte de l'article.")
    print()
    print("  2. N_fin / N_final. La these declare  N_fin <- N - eval  puis")
    print("     teste  N_final > N_inner : deux noms pour la meme variable.")
    print("     Le pseudocode utilise N_final partout.")
    print()
    print("  Pour revenir a la version de la these, editez LINES en tete")
    print("  de ce script.")


if __name__ == '__main__':
    main()