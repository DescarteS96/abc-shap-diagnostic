"""
make_article_tables.py
======================
Construit les tableaux 1, 2 et 3 de l'article, au format Word (.docx),
directement a partir des CSV bruts audites.

  Table 1 — Ecart de performance : correlation de Spearman moyenne par
            cellule (dataset x budget), cinq methodes, tests de Wilcoxon
            apparies contre ABC-SHAP.
            Entree : les CSV de la grille principale (un par cellule).

  Table 2 — Effet du ratio d'allocation A0..A3 : allocation, Spearman,
            Kendall, top-3, ratio MSE relatif a B2.
            Entree : le CSV consolide de l'experience d'allocation.

  Table 3 — Validite du signal d'aptitude : correlations de Pearson,
            Spearman et Kendall entre J(pi) et la vraie stabilite
            ordinale, par dataset puis en agrege.
            Entree : j_pi_true_correlation.csv.

CONVENTIONS APPLIQUEES (identiques dans les trois tableaux)
  - Ridge sur California est ordinalement degenere : modele lineaire sur
    une tache de regression, l'ordre des attributions est invariant a la
    politique d'echantillonnage et les cinq methodes retournent une
    correlation de Spearman de 1.0000 avec variance nulle. Le test
    apparie y est indefini. Cette configuration est donc EXCLUE des
    metriques de rang et CONSERVEE pour le ratio MSE.
  - Tests de Wilcoxon signes sur les moyennes par instance.
  - Correction de Bonferroni a l'interieur de chaque cellule :
    alpha = 0.05 / 4 = 0.0125 (quatre baselines).
  - Le signe de la difference est rapporte : un ecart significatif en
    faveur d'ABC-SHAP n'est pas note comme un ecart defavorable.

CONTROLES BLOQUANTS
  Le script refuse de produire un tableau si un CSV de la grille
  principale contient plus d'un budget d'exploration : cela signale un
  fichier produit sous deux formules d'allocation differentes, qui ne
  peut pas etre agrege.

Usage :
  python make_article_tables.py \
      --main results_california_N1000.csv results_california_N5000.csv \
             results_adult_N1000.csv results_adult_N5000.csv \
             results_credit_N1000.csv results_credit_N5000.csv \
      --allocation results_allocation_consolidated_california_N5000.csv \
      --fitness j_pi_true_correlation.csv \
      --outdir tables_article

Chaque bloc est optionnel : passer seulement --main ne produit que la
Table 1.
"""

import argparse
import os
import sys

import numpy as np
import pandas as pd
from scipy.stats import wilcoxon, pearsonr, spearmanr, kendalltau

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERREUR : python-docx n'est pas installe.")
    print("Lancez : pip install python-docx")
    sys.exit(1)

# ── Conventions ───────────────────────────────────────────────────────────
ALPHA = 0.0125                      # 0.05 / 4 baselines, Bonferroni par cellule
BASELINES = ['b1_uniform', 'b2_stratified', 'b3_is', 'b4_antithetic']
BASELINE_LABEL = {'b1_uniform': 'B1', 'b2_stratified': 'B2',
                  'b3_is': 'B3', 'b4_antithetic': 'B4'}
DATASET_LABEL = {'california': 'California', 'adult': 'Adult',
                 'credit': 'Credit Default'}
DATASET_ORDER = ['california', 'adult', 'credit']
MODEL_ORDER = ['ridge', 'rf', 'mlp']
MODEL_LABEL = {'ridge': 'Ridge / Log.', 'rf': 'RF', 'mlp': 'MLP'}

# (dataset, model) exclus des metriques de rang — voir en-tete
RANK_DEGENERATE = {('california', 'ridge')}

ALLOC_ORDER = ['A0', 'A3', 'A2', 'A1']   # par part d'estimation croissante


def log(m):
    print(f"[tables] {m}")


def die(m):
    print(f"[tables] ERREUR — {m}")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════════════
# HELPERS WORD
# ══════════════════════════════════════════════════════════════════════════

def new_document():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(10)
    return doc


def add_caption(doc, text):
    """Legende de tableau : au-dessus, alignee a gauche, numero en gras."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    parts = text.split('. ', 1)
    r = p.add_run(parts[0] + '.')
    r.bold = True
    r.font.size = Pt(10)
    if len(parts) > 1:
        r2 = p.add_run(' ' + parts[1])
        r2.italic = True
        r2.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(8)
    return p


def add_table(doc, header, rows, widths_cm=None):
    """Tableau simple, style Table Grid, en-tete en gras."""
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run('' if v is None else str(v))
            r.font.size = Pt(9)
    if widths_cm:
        # Layout fixe : sans cela Word et LibreOffice repartissent les
        # colonnes uniformement et ignorent les largeurs demandees.
        tblPr = t._tbl.tblPr
        for el in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(el)
        lay = OxmlElement('w:tblLayout')
        lay.set(qn('w:type'), 'fixed')
        tblPr.append(lay)
        # tblGrid : c'est cette grille que Word et LibreOffice lisent en
        # priorite ; les largeurs de cellule seules sont ignorees.
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            t._tbl.remove(grid)
        grid = OxmlElement('w:tblGrid')
        for w in widths_cm:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(round(Cm(w).twips))))
            grid.append(gc)
        t._tbl.insert(list(t._tbl).index(tblPr) + 1, grid)
        for j, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


# ══════════════════════════════════════════════════════════════════════════
# TABLE 1 — ecart de performance
# ══════════════════════════════════════════════════════════════════════════

def load_main_cells(paths):
    """Charge et valide chaque CSV de la grille principale."""
    cells = []
    for p in paths:
        if not os.path.exists(p):
            die(f"fichier introuvable : {p}")
        d = pd.read_csv(p)
        name = os.path.basename(p)

        for c in ['dataset', 'model', 'method', 'instance_idx',
                  'spearman_mean', 'n_exploration', 'n_estimation', 'n_budget']:
            if c not in d.columns:
                die(f"{name} : colonne '{c}' absente.")

        if d['dataset'].nunique() != 1 or d['n_budget'].nunique() != 1:
            die(f"{name} : plusieurs datasets ou budgets dans le meme fichier.")

        abc = d[d['method'] == 'abc_shap']
        expl = sorted(set(abc['n_exploration'].astype(int)))
        if len(expl) != 1:
            die(f"{name} : {len(expl)} budgets d'exploration differents "
                f"({expl}). Ce fichier melange deux regles d'allocation et "
                f"ne peut pas etre agrege. Relancez la cellule entierement.")

        dup = d.duplicated(['model', 'method', 'instance_idx']).sum()
        if dup:
            die(f"{name} : {dup} lignes dupliquees sur "
                f"(model, method, instance_idx).")

        ds = str(d['dataset'].iloc[0])
        N = int(d['n_budget'].iloc[0])
        M = int(d['M'].iloc[0])
        est = int(abc['n_estimation'].iloc[0])
        log(f"{name:38s} {len(d):4d} l | M={M:2d} N={N:5d} | "
            f"expl={expl[0]:5d} est={est:5d} ({100.0*est/N:4.1f}%)")
        cells.append(dict(df=d, dataset=ds, N=N, M=M,
                          expl=expl[0], est=est, name=name))
    cells.sort(key=lambda c: (DATASET_ORDER.index(c['dataset'])
                              if c['dataset'] in DATASET_ORDER else 99, c['N']))
    return cells


def mark(p, med):
    """Marqueur de significativite, signe."""
    if p != p:
        return 'n/a'
    if p > ALPHA:
        return ''
    s = '\u207a' if med > 0 else '\u207b'
    return s * 2 if p < 0.001 else s


def build_table1(cells):
    rows = []
    audit = []
    for c in cells:
        d = c['df']
        ds, N = c['dataset'], c['N']
        keep = [m for m in MODEL_ORDER
                if m in set(d['model']) and (ds, m) not in RANK_DEGENERATE]
        sub = d[d['model'].isin(keep)]
        w = sub.pivot_table(index=['model', 'instance_idx'],
                            columns='method', values='spearman_mean')
        n = len(w)
        row = [DATASET_LABEL.get(ds, ds), str(c['M']), f"{N:,}".replace(',', ' '),
               f"{100.0 * c['est'] / N:.1f}", str(n),
               f"{w['abc_shap'].mean():.3f} ± {w['abc_shap'].std(ddof=1):.3f}"]
        for b in BASELINES:
            st, p = wilcoxon(w['abc_shap'], w[b])
            med = float(np.median(w['abc_shap'] - w[b]))
            row.append(f"{w[b].mean():.3f} {mark(p, med)}".strip())
            audit.append((ds, N, BASELINE_LABEL[b], med, p, n))
        rows.append(row)
    header = ['Dataset', 'M', 'N', 'Est. %', 'n',
              'ABC-SHAP', 'B1', 'B2', 'B3', 'B4']
    return header, rows, audit


def write_table1(cells, outdir):
    header, rows, audit = build_table1(cells)
    doc = new_document()
    add_caption(doc,
        "Table 1. Mean Spearman correlation with the high-budget reference, "
        "by dataset and evaluation budget, and paired Wilcoxon tests against "
        "ABC-SHAP.")
    add_table(doc, header, rows,
              widths_cm=[2.5, 0.8, 1.3, 1.2, 0.7, 2.5, 1.7, 1.7, 1.7, 1.7])
    add_note(doc,
        "Paired Wilcoxon signed-rank tests on per-instance mean Spearman "
        "correlations, Bonferroni-corrected within each cell "
        f"(\u03b1 = 0.05/4 = {ALPHA}). Superscripts denote the direction of a "
        "significant difference: \u207a\u207a ABC-SHAP significantly higher "
        "(p < 0.001), \u207a higher (p \u2264 0.0125), \u207b\u207b "
        "significantly lower (p < 0.001), \u207b lower (p \u2264 0.0125); no "
        "mark indicates no significant difference. Ridge on California "
        "Housing is excluded from the rank-based comparison: as a linear "
        "model on a regression task its attribution ordering is invariant to "
        "the sampling policy, all five methods return a Spearman correlation "
        "of exactly 1.0000 with zero variance, and the paired test is "
        "undefined; the configuration is retained for the accuracy "
        "comparison. K = 30 independent repetitions per instance and method.")
    p = os.path.join(outdir, 'table1_performance_gap.docx')
    doc.save(p)
    log(f"ecrit : {p}")

    # ── Audit console ─────────────────────────────────────────────────────
    print()
    log("Table 1 — detail des tests (verification) :")
    print(f"  {'cellule':<26}{'baseline':>9}{'mediane':>10}{'p':>10}   verdict")
    tot = sig_neg = sig_pos = 0
    for ds, N, b, med, p_, n in audit:
        tot += 1
        v = mark(p_, med)
        if v.startswith('\u207b'):
            sig_neg += 1
        elif v.startswith('\u207a'):
            sig_pos += 1
        print(f"  {DATASET_LABEL.get(ds,ds)+' N='+str(N):<26}{b:>9}"
              f"{med:>+10.4f}{p_:>10.5f}   {v or 'ns'}")
    print()
    log(f"  total = {tot} | ABC significativement inferieur = {sig_neg} | "
        f"superieur = {sig_pos} | non significatif = {tot-sig_neg-sig_pos}")
    b23 = [(m, p_) for ds, N, b, m, p_, n in audit if b in ('B2', 'B3')]
    ok = sum(1 for m, p_ in b23 if p_ <= ALPHA and m < 0)
    log(f"  vs B2/B3 : {ok}/{len(b23)} significativement defavorables")


# ══════════════════════════════════════════════════════════════════════════
# TABLE 2 — effet du ratio d'allocation
# ══════════════════════════════════════════════════════════════════════════

def write_table2(path, outdir):
    if not os.path.exists(path):
        die(f"fichier introuvable : {path}")
    d = pd.read_csv(path)
    for c in ['allocation', 'method', 'model', 'instance_idx', 'dataset',
              'spearman_mean', 'kendall_mean', 'top3_mean', 'mse_ratio_b2',
              'n_exploration', 'n_estimation', 'n_budget']:
        if c not in d.columns:
            die(f"{os.path.basename(path)} : colonne '{c}' absente.")

    ds = str(d['dataset'].iloc[0])
    N = int(d['n_budget'].iloc[0])
    a = d[(d['method'] == 'abc_shap') &
          (~d.apply(lambda r: (r['dataset'], r['model']) in RANK_DEGENERATE,
                    axis=1))]
    present = [x for x in ALLOC_ORDER if x in set(a['allocation'].astype(str))]
    if len(present) < 2:
        die("moins de deux allocations : tableau sans objet.")
    log(f"Table 2 : allocations presentes = {present}")

    rows = []
    for al in present:
        s = a[a['allocation'] == al]
        est = int(s['n_estimation'].iloc[0])
        rows.append([
            al,
            f"{int(s['n_exploration'].iloc[0]):,}".replace(',', ' '),
            f"{est:,}".replace(',', ' '),
            f"{100.0 * est / N:.1f}",
            f"{s['spearman_mean'].mean():.3f} ± {s['spearman_mean'].std(ddof=1):.3f}",
            f"{s['kendall_mean'].mean():.3f}",
            f"{s['top3_mean'].mean():.3f}",
            f"{s['mse_ratio_b2'].mean():.1f}",
            str(len(s)),
        ])

    doc = new_document()
    add_caption(doc,
        "Table 2. Effect of the exploration/estimation budget split on "
        "ABC-SHAP at a fixed total evaluation budget.")
    add_table(doc,
              ['Allocation', 'Explor.', 'Estim.', 'Est. %',
               'Spearman', 'Kendall', 'Top-3', 'MSE / B2', 'n'],
              rows,
              widths_cm=[1.8, 1.8, 1.8, 1.5, 2.6, 1.5, 1.3, 2.0, 0.7])
    nfmt = f"{N:,}".replace(',', '\u2009')
    add_note(doc,
        f"{DATASET_LABEL.get(ds, ds)}, N = {nfmt}, "
        "K = 30 repetitions per instance. Only the allocation ratio varies: "
        "dataset, model families, test instances, background sample, random "
        "seed, ABC hyperparameters and the high-budget reference are identical "
        "across allocations. A0 is the initial allocation; A1, A2 and A3 "
        "target 25%, 50% and 75% of the budget for exploration. Because the "
        "exploration cost is a multiple of N_inner, realized shares differ "
        "slightly from the nominal targets. Ridge is excluded (see Table 1).")
    p = os.path.join(outdir, 'table2_allocation_effect.docx')
    doc.save(p)
    log(f"ecrit : {p}")

    # ── Tests apparies, en console ────────────────────────────────────────
    piv = a.pivot_table(index=['model', 'instance_idx'],
                        columns='allocation', values='spearman_mean')
    bal = [x for x in ['A1', 'A2', 'A3'] if x in present]
    print()
    if 'A0' in present and bal:
        alpha_f = 0.05 / len(bal)
        log(f"Table 2 — A0 vs allocations equilibrees (alpha = {alpha_f:.4f}) :")
        for b in bal:
            dd = piv[b] - piv['A0']
            st, p_ = wilcoxon(piv[b], piv['A0'])
            print(f"    {b} vs A0  mediane={np.median(dd):+.4f}  "
                  f"mieux={int((dd>0).sum())}/{len(dd)}  p={p_:.6f}  "
                  f"{'SIG' if p_ <= alpha_f else 'ns'}")
    if len(bal) >= 2:
        import itertools
        alpha_g = 0.05 / len(list(itertools.combinations(bal, 2)))
        log(f"Table 2 — entre allocations equilibrees (alpha = {alpha_g:.4f}) :")
        for x, y in itertools.combinations(bal, 2):
            dd = piv[y] - piv[x]
            st, p_ = wilcoxon(piv[y], piv[x])
            print(f"    {y} vs {x}  mediane={np.median(dd):+.4f}  "
                  f"mieux={int((dd>0).sum())}/{len(dd)}  p={p_:.6f}  "
                  f"{'SIG' if p_ <= alpha_g else 'ns'}")


# ══════════════════════════════════════════════════════════════════════════
# TABLE 3 — validite du signal d'aptitude
# ══════════════════════════════════════════════════════════════════════════

TARGETS = [
    ('true_stability_spearman', 'J(\u03c0) vs Spearman'),
    ('true_stability_kendall',  'J(\u03c0) vs Kendall'),
    ('true_stability_top3',     'J(\u03c0) vs top-3'),
]


def triplet(x, y):
    x = np.asarray(x, float)
    y = np.asarray(y, float)
    ok = ~(np.isnan(x) | np.isnan(y))
    x, y = x[ok], y[ok]
    if len(x) < 3 or np.std(x) == 0 or np.std(y) == 0:
        return (np.nan,) * 6
    rp, pp = pearsonr(x, y)
    rs, ps = spearmanr(x, y)
    rk, pk = kendalltau(x, y)
    return float(rp), float(pp), float(rs), float(ps), float(rk), float(pk)


def write_table3(path, outdir):
    if not os.path.exists(path):
        die(f"fichier introuvable : {path}")
    d = pd.read_csv(path)
    need = ['dataset', 'J_pi'] + [c for c, _ in TARGETS]
    for c in need:
        if c not in d.columns:
            die(f"{os.path.basename(path)} : colonne '{c}' absente. "
                f"Relancez test_j_pi_true_correlation.py dans sa version "
                f"produisant true_stability_top3.")
    log(f"Table 3 : {len(d)} politiques | "
        f"{d.groupby('dataset').size().to_dict()}")

    rows = []
    for ds in DATASET_ORDER:
        sub = d[d['dataset'] == ds]
        if len(sub) < 5:
            continue
        M = int(sub['M'].iloc[0]) if 'M' in sub.columns else None
        scope = DATASET_LABEL.get(ds, ds) + (f" (M = {M})" if M else '')
        for col, lab in TARGETS:
            rp, pp, rs, ps, rk, pk = triplet(sub['J_pi'], sub[col])
            rows.append([scope, lab, str(len(sub)),
                         f"{rp:+.3f}", f"{pp:.3f}",
                         f"{rs:+.3f}", f"{ps:.3f}",
                         f"{rk:+.3f}", f"{pk:.3f}"])
            scope = ''

    # Agregat : tout centre-reduit a l'interieur de chaque dataset, les
    # echelles n'etant pas comparables entre dimensionnalites.
    z = d.copy()
    for c in ['J_pi'] + [c for c, _ in TARGETS]:
        z[c + '_z'] = z.groupby('dataset')[c].transform(
            lambda s: (s - s.mean()) / s.std(ddof=0) if s.std(ddof=0) > 0 else 0.0)
    scope = 'Pooled (within-dataset z-scores)'
    for col, lab in TARGETS:
        rp, pp, rs, ps, rk, pk = triplet(z['J_pi_z'], z[col + '_z'])
        rows.append([scope, lab, str(len(z)),
                     f"{rp:+.3f}", f"{pp:.3f}",
                     f"{rs:+.3f}", f"{ps:.3f}",
                     f"{rk:+.3f}", f"{pk:.3f}"])
        scope = ''

    doc = new_document()
    add_caption(doc,
        "Table 3. Relationship between the internal split-half fitness "
        "J(\u03c0) and true inter-run ordinal stability, measured "
        "independently under the same fixed policy.")
    add_table(doc,
              ['Scope', 'Relation', 'n', 'Pearson', 'p',
               'Spearman', 'p', 'Kendall', 'p'],
              rows,
              widths_cm=[3.1, 2.7, 0.7, 1.6, 1.2, 1.6, 1.2, 1.6, 1.2])
    add_note(doc,
        "Thirty candidate policies per dataset (ten drawn uniformly at random "
        "on the simplex, twenty visited during an actual ABC search). True "
        "stability is the mean pairwise agreement across K = 30 independent "
        "full estimations under the same fixed policy. The pooled rows "
        "standardize J(\u03c0) and the stability measures within each dataset, "
        "since their scales are not comparable across dimensionalities.")
    p = os.path.join(outdir, 'table3_fitness_validity.docx')
    doc.save(p)
    log(f"ecrit : {p}")

    print()
    log("Table 3 — synthese :")
    allp = [float(r[4]) for r in rows] + [float(r[6]) for r in rows] + \
           [float(r[8]) for r in rows]
    allr = [abs(float(r[3])) for r in rows] + [abs(float(r[5])) for r in rows] + \
           [abs(float(r[7])) for r in rows]
    log(f"  {sum(1 for x in allp if x <= 0.05)}/{len(allp)} coefficients "
        f"atteignent p <= 0.05")
    log(f"  coefficient absolu maximal = {max(allr):.4f}")



# ══════════════════════════════════════════════════════════════════════════
# TABLE 4 — controle Dirichlet
# ══════════════════════════════════════════════════════════════════════════

def write_table4(path, outdir, keep_partial=False):
    """Controle : la recherche par colonie fait-elle mieux qu'un tirage
    aleatoire de politique, a budget identique et mecanisme identique ?"""
    if not os.path.exists(path):
        die(f"fichier introuvable : {path}")
    d = pd.read_csv(path)
    for c in ['dataset', 'model', 'instance_idx', 'method',
              'spearman_mean', 'mse_mean']:
        if c not in d.columns:
            die(f"{os.path.basename(path)} : colonne '{c}' absente.")

    log(f"Table 4 : {len(d)} lignes | "
        f"{sorted(set(d['dataset']))} | methodes {sorted(set(d['method']))}")

    # ── Blocs incomplets ────────────────────────────────────────────────
    # Un run interrompu laisse des blocs partiels. Les agreger produirait
    # un effectif desequilibre et une conclusion qui depend du point
    # d'arret, pas des donnees. Ils sont ecartes par defaut.
    counts = (d[d['method'] == 'abc_corrected']
              .groupby(['dataset', 'model'])['instance_idx'].nunique())
    if len(counts) == 0:
        die("aucune ligne 'abc_corrected' dans le fichier.")
    n_full = int(counts.max())
    n_models = int(counts.reset_index().groupby('dataset')['model'].nunique().max())

    if not keep_partial:
        # Un dataset n'est retenu que si TOUTES ses familles de modeles sont
        # completes. Retenir un dataset represente par une seule famille
        # reviendrait a melanger des perimetres differents dans la ligne
        # agregee, et la conclusion dependrait du moment ou le run a ete
        # interrompu plutot que des donnees.
        complete = counts[counts == n_full].reset_index()
        ok = (complete.groupby('dataset')['model'].nunique() == n_models)
        keep_ds = set(ok[ok].index)
        drop_ds = sorted(set(counts.reset_index()['dataset']) - keep_ds)
        if drop_ds:
            log(f"  datasets ecartes (couverture incomplete, "
                f"{n_models} familles de modeles x {n_full} instances "
                f"attendues) :")
            for ds_ in drop_ds:
                det = counts.loc[ds_]
                log(f"    {ds_} : " + ", ".join(
                    f"{m_} {int(c)}/{n_full}" for m_, c in det.items()))
            log("  utilisez --dirichlet-keep-partial pour les conserver.")
            d = d[d['dataset'].isin(keep_ds)]
        log(f"  perimetre retenu : {sorted(keep_ds)}")
    else:
        log("  ATTENTION : --dirichlet-keep-partial actif, aucun filtrage.")

    rows = []
    deltas_all = []
    for ds in DATASET_ORDER:
        sub = d[d['dataset'] == ds]
        if sub.empty:
            continue
        for m in MODEL_ORDER:
            s = sub[sub['model'] == m]
            if s.empty:
                continue
            w = s.pivot_table(index='instance_idx', columns='method',
                              values='spearman_mean')
            if 'abc_corrected' not in w.columns or 'dirichlet_strong' not in w.columns:
                continue
            diff = (w['abc_corrected'] - w['dirichlet_strong']).dropna()
            degenerate = bool(np.allclose(diff.values, 0.0))
            if degenerate:
                rows.append([DATASET_LABEL.get(ds, ds), MODEL_LABEL.get(m, m),
                             str(len(diff)), '0.000', '0.000',
                             f"0/{len(diff)}", 'undef.'])
                continue
            st, pv = wilcoxon(w['abc_corrected'], w['dirichlet_strong'])
            deltas_all.extend(diff.values.tolist())
            rows.append([
                DATASET_LABEL.get(ds, ds), MODEL_LABEL.get(m, m), str(len(diff)),
                f"{diff.mean():+.4f}", f"{np.median(diff):+.4f}",
                f"{int((diff > 0).sum())}/{len(diff)}", f"{pv:.3f}",
            ])

    if not deltas_all:
        die("aucune comparaison informative : verifiez le fichier.")
    a = np.array(deltas_all)
    st, pv = wilcoxon(a)
    rng = np.random.RandomState(42)
    boot = a[rng.randint(0, len(a), (5000, len(a)))].mean(axis=1)
    lo, hi = np.percentile(boot, 2.5), np.percentile(boot, 97.5)
    rows.append(['Pooled (informative)', '', str(len(a)),
                 f"{a.mean():+.4f}", f"{np.median(a):+.4f}",
                 f"{int((a > 0).sum())}/{len(a)}", f"{pv:.3f}"])

    doc = new_document()
    add_caption(doc,
        "Table 4. Control experiment: adaptive colony search versus a single "
        "random policy draw, at identical total budget.")
    add_table(doc,
              ['Dataset', 'Model', 'n', 'Mean \u0394\u03c1', 'Median \u0394\u03c1',
               'ABC higher', 'p'],
              rows,
              widths_cm=[3.2, 1.9, 0.8, 2.1, 2.2, 2.0, 1.3])
    add_note(doc,
        "\u0394\u03c1 is the paired difference in mean Spearman correlation "
        "between ABC-SHAP and a single Dirichlet policy draw spending the full "
        "budget on estimation, computed per test instance "
        f"(N = 5000, K = 30). Paired Wilcoxon signed-rank tests. Bootstrap 95% "
        f"confidence interval on the pooled mean: [{lo:+.4f}, {hi:+.4f}]. "
        "Ridge on California Housing yields differences that are exactly zero "
        "for every instance and is excluded from the pooled row; the test is "
        "undefined there (see Table 1). Both methods draw coalitions by the "
        "same multinomial mechanism and differ only in how the sampling policy "
        "is obtained. Only blocks with a complete set of test instances are "
        "reported.")
    pth = os.path.join(outdir, 'table4_dirichlet_control.docx')
    doc.save(pth)
    log(f"ecrit : {pth}")

    print()
    log("Table 4 — synthese :")
    log(f"  pooled n={len(a)}  moyenne={a.mean():+.4f}  mediane={np.median(a):+.4f}"
        f"  ABC>Dir={int((a>0).sum())}/{len(a)}  p={pv:.4f}")
    log(f"  IC bootstrap 95% : [{lo:+.4f}, {hi:+.4f}]")


# ══════════════════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--main', nargs='+', default=None,
                    help="CSV de la grille principale (un par cellule)")
    ap.add_argument('--allocation', default=None,
                    help="CSV consolide de l'experience d'allocation")
    ap.add_argument('--fitness', default=None,
                    help="j_pi_true_correlation.csv")
    ap.add_argument('--dirichlet', default=None,
                    help="abc_vs_dirichlet_results.csv")
    ap.add_argument('--dirichlet-keep-partial', action='store_true',
                    help="conserver les blocs incomplets du controle Dirichlet")
    ap.add_argument('--outdir', default='tables_article')
    args = ap.parse_args()

    if not any([args.main, args.allocation, args.fitness, args.dirichlet]):
        die("aucune entree : passez au moins --main, --allocation, "
            "--fitness ou --dirichlet.")
    os.makedirs(args.outdir, exist_ok=True)

    if args.main:
        print()
        log("=== TABLE 1 ===")
        write_table1(load_main_cells(args.main), args.outdir)
    if args.allocation:
        print()
        log("=== TABLE 2 ===")
        write_table2(args.allocation, args.outdir)
    if args.fitness:
        print()
        log("=== TABLE 3 ===")
        write_table3(args.fitness, args.outdir)
    if args.dirichlet:
        print()
        log("=== TABLE 4 ===")
        write_table4(args.dirichlet, args.outdir, args.dirichlet_keep_partial)

    print()
    log(f"Termine. Fichiers dans : {os.path.abspath(args.outdir)}")


if __name__ == '__main__':
    main()